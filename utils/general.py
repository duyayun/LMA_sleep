
import numpy as np
import mne
import pandas as pd
import polars as pl
from scipy import signal
import docx
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

def file_to_lazy_frame(filename):
    extension = filename.split('.')[-1]
    if extension == 'csv':
        ret = pl.read_csv(filename, columns=['t', 'x', 'y', 'z'], use_pyarrow=True).select([
            pl.col('ts').cast(pl.Datetime), 
            pl.col('x').cast(pl.Int16),
            pl.col('y').cast(pl.Int16),
            pl.col('z').cast(pl.Int16)
            ]).interpolate().filter(pl.col('ts').is_not_null()).filter(pl.col('z').is_not_null())
        
    elif extension == 'parquet':
        ret = pl.read_parquet(filename, columns=['t', 'x', 'y', 'z'], use_pyarrow=True).select([
            pl.col('t').cast(pl.Datetime), 
            pl.col('x').cast(pl.Int16),
            pl.col('y').cast(pl.Int16),
            pl.col('z').cast(pl.Int16)
            ]).interpolate().filter(pl.col('t').is_not_null()).filter(pl.col('z').is_not_null())
    elif extension == 'gz':
        print('gz format is not yet supported')

    else:
        print(f'bad file format: {extension}')    
    
    
    return ret



def shannon_energy(x):
    """Implementation of shannon energy

    Args:
        x (_type_): input signal
    """
    x_env = -x**2 * np.log(x**2)
    return(x_env)


def bandpass_filt(sig,cutoff,fs,mode):
    """ bandpass filtering

    Args:
        sig (_type_): input signal. 
        cutoff (_type_): cutoff requency
        fs (_type_): signal frequency
        mode (_type_): choice between 'lowpass' 'highpass' and 'bandpass'

    Returns:
        _type_: filtered signal
    """
    nyq = 0.5*fs
    wn = cutoff/nyq
    sos = signal.butter(4, wn, btype=mode,output='sos')
    return pd.DataFrame(data=signal.sosfiltfilt(sos,sig,axis=0),columns=sig.columns,index=sig.index)


def read_edf(path):
    """reads edf file into dataframe, and upsample to 1kHz
    

    Args:
        path (_type_): _description_

    Returns:
        _type_: _description_
    """
    raw = mne.io.read_raw_edf(path)
    start_time = raw.info['meas_date'].replace(tzinfo=None)
    df = raw.to_data_frame()
    df['time'] = pd.to_timedelta(df['time'], unit='s') + start_time
    df.set_index('time', inplace=True)
    df = df[['PR', 'RR']]
    df_resampled = df.resample('1ms').asfreq()
    df_upsampled = df_resampled.interpolate(method='linear')
    
    df_upsampled['PR'] *= 1e-6

    return df_upsampled


def label_dataframe(X, doc):
    doc = docx.Document(doc)
    start_date_string = None
    for s in doc.paragraphs:
        if 'Date' in s.text:
            start_date_string = s.text

    assert start_date_string, "Date is not found in docx"

    format_string = "Date: %m/%d/%Y %I:%M:%S %p"

    doc_start_time = datetime.strptime(start_date_string, format_string)
    for table in doc.tables:
        # Extract the contents of each cell in the table
        data = [[cell.text for cell in row.cells] for row in table.rows]
        
        # Create a pandas DataFrame from the extracted data
        df = pd.DataFrame(data[1:], columns=data[0])

    stage_values = df['Stage'].values.T.flatten()
    epoch_values = df['Epoch'].values.T.flatten()

    label_df = pd.DataFrame()
    label_df['Stage'] = stage_values
    label_df['Epoch'] = epoch_values

    doc_start_time = doc_start_time
    increment = timedelta(seconds=30)
    num_periods = len(label_df)
    time_index = pd.date_range(start=doc_start_time, periods=num_periods, freq=increment)
    label_df.set_index(time_index, inplace=True)
    final_df = pd.concat([X, label_df], axis=1)
    final_df['Stage'] = final_df['Stage'].fillna(method='ffill')
    final_df['Epoch'] = final_df['Epoch'].fillna(method='ffill')
    # label_index = list(set(final_df['Stage']))
    label_map={
        '<Stage>Wake':0,
        '<Stage>NREM 1':1,
        '<Stage>NREM 2':2,
        '<Stage>NREM 3':3,
        '<Stage>REM':4,
        '<Stage>UNS':5,
        '<Stage>-':6,
        '<Stage>Unscored':7
    }
    # print(label_map)
    final_df['Stage'].map(label_map)
    final_df['QuantizedStage'] = final_df['Stage'].map(label_map)

    return final_df



def plot_labeled_dataframe(df):
    # reversed_label_map = {value: key for key, value in label_map.items()}

    fig, ax = plt.subplots(2, figsize=(30, 10))
    color_map = {'<Stage>Wake':'red', '<Stage>NREM 1':'green', '<Stage>NREM 2':'blue', '<Stage>NREM 3': 'orange', '<Stage>REM':'black', '<Stage>-': 'lavender', '<Stage>UNS':'purple'}
    legend_patches = []
    # for i in range(len(stage_index[:-1])):
    #     ax.axvspan(stage_index[i], stage_index[i + 1], facecolor=color_map[final_df['QuantizedStage'][stage_index[i]]], alpha=0.5)
    for i, label in enumerate(df['Stage']):
        if i == 0:
            continue
        if df['Stage'][i] != df['Stage'][i - 1]:
            ax[0].axvspan(df.index[i - 1], df.index[i], facecolor=color_map[df['label'][i - 1]], alpha=0.5)
            ax[1].axvspan(df.index[i - 1], df.index[i], facecolor=color_map[df['label'][i - 1]], alpha=0.5)

    for k, v in color_map.items():
        legend_patches.append(mpatches.Patch(color=v, label=k, alpha=0.5))
    ax[0].plot(df.index, df['PR'], label='Heart Rate', color='black', linewidth=2)
    ax[0].plot(df.index, df['ma_hr'], label='MA Heart Rate', color='cyan', linewidth=1)
    ax[0].set_xlabel('Time(Month-Day Hour)', fontsize=24)
    ax[0].set_ylabel('Heart Rate', fontsize=24)
    ax[0].set_title('Heart Rate with Sleep Stage Labels', fontsize=36)
    ax[0].set_ylim(0, 120)
    # Add the legend with custom patches
    ax[0].legend(handles=list(set(legend_patches + [mpatches.Patch(color='black', label='Heart Rate'),mpatches.Patch(color='cyan', label='MA Heart Rate')])), loc='upper left', bbox_to_anchor=(1, 1), fontsize=24)

    
    ax[1].plot(df.index, df['RR'], label='EDF Respiration Rate', color='black', linewidth=2)
    ax[1].plot(df.index, df['ma_rr'], label='MA Respiration Rate', color='cyan', linewidth=1)
    ax[1].set_xlabel('Time(Month-Day Hour)', fontsize=24)
    ax[1].set_ylabel('Heart Rate', fontsize=24)
    ax[1].set_ylim(0, 120)
    ax[1].set_title('Respiration Rate with Sleep Stage Labels', fontsize=36)
    # ax[1].set_ylim(0, 40)
    ax[1].legend(handles=list(set(legend_patches + [mpatches.Patch(color='black', label='Heart Rate'),mpatches.Patch(color='cyan', label='MA Heart Rate')])), loc='upper left', bbox_to_anchor=(1, 1), fontsize=24)


    plt.show()


def plot_comparison(df):
    reversed_label_map = {value: key for key, value in label_map.items()}

    fig, ax = plt.subplots(2, figsize=(30, 20))
    for a in ax:
        a.tick_params(axis='both', labelsize=20)
    color_map = {0:'red', 1:'green', 2:'blue', 3: 'orange', 4:'black', 5: 'lavender', 6:'purple', 7:'yellow'}
    legend_patches = []
    for i in range(len(stage_index[:-1])):
        ax[0].axvspan(stage_index[i], stage_index[i + 1], facecolor=color_map[final_df['QuantizedStage'][stage_index[i]]], alpha=0.5)
        ax[1].axvspan(stage_index[i], stage_index[i + 1], facecolor=color_map[final_df['QuantizedStage'][stage_index[i]]], alpha=0.5)
    for k, v in reversed_label_map.items():
        legend_patches.append(mpatches.Patch(color=color_map[k], label=v, alpha=0.3))
    ax[0].plot(df.index, df['PR'], label='Heart Rate', color='black', linewidth=2)
    ax[0].plot(df.index, df['ma_hr'], label='MA Heart Rate', color='cyan', linewidth=1)
    ax[0].set_xlabel('Time(Month-Day Hour)', fontsize=24)
    ax[0].set_ylabel('Heart Rate', fontsize=24)
    ax[0].set_title('Heart Rate with Sleep Stage Labels (Current Pipeline)', fontsize=36)
    ax[0].set_ylim(0, 120)
    # Add the legend with custom patches
    ax[0].legend(handles=list(set(legend_patches + [mpatches.Patch(color='black', label='Heart Rate'),mpatches.Patch(color='cyan', label='MA Heart Rate')])), loc='upper left', bbox_to_anchor=(1, 1), fontsize=24)

    
    ax[1].plot(df.index, df['PR'], label='EDF Heart Rate', color='black', linewidth=2)
    ax[1].plot(df.index, df['a_hr_ma'], label='MA Heart Rate', color='cyan', linewidth=1)
    ax[1].set_xlabel('Time(Month-Day Hour)', fontsize=24)
    ax[1].set_ylabel('Heart Rate', fontsize=24)
    ax[1].set_ylim(0, 120)
    ax[1].set_title('Heart Rate with Sleep Stage Labels (Andreas\'s Pipeline)', fontsize=36)
    # ax[1].set_ylim(0, 40)
    ax[1].legend(handles=list(set(legend_patches + [mpatches.Patch(color='black', label='Heart Rate'),mpatches.Patch(color='cyan', label='MA Heart Rate')])), loc='upper left', bbox_to_anchor=(1, 1), fontsize=24)


    plt.show()

